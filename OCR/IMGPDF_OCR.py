import os
import cv2
import numpy as np
import pytesseract
import pandas as pd
from sklearn.cluster import AgglomerativeClustering
from pytesseract import Output
from pdf2image import convert_from_path
import PyPDF2
from docx import Document
import tkinter as tk
from tkinter import filedialog, ttk, messagebox

# Note: This script requires the following installations:
# pip install opencv-python pytesseract pandas openpyxl pdf2image PyPDF2 python-docx scikit-learn
# Additionally, you need to have Tesseract OCR installed on your system[](https://github.com/tesseract-ocr/tesseract)
# and Poppler for pdf2image[](https://pdf2image.readthedocs.io/en/latest/installation.html)

def get_pdf_page_count(pdf_path):
    with open(pdf_path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        return len(reader.pages)

def populate_page_ranges(num_pages):
    ranges = ['All']
    if num_pages > 1:
        batch_size = 5 if num_pages <= 20 else 20 if num_pages <= 100 else 50
        for i in range(1, num_pages + 1, batch_size):
            end = min(i + batch_size - 1, num_pages)
            ranges.append(f"{i}-{end}")
    return ranges

def select_file():
    file_path = filedialog.askopenfilename(filetypes=[("PDF or Image files", "*.pdf *.jpg *.jpeg *.png")])
    if file_path:
        entry_file_path.set(file_path)
        if file_path.lower().endswith('.pdf'):
            num_pages = get_pdf_page_count(file_path)
            page_ranges = populate_page_ranges(num_pages)
            combo_pages['values'] = page_ranges
            combo_pages.current(0)
        else:
            combo_pages['values'] = ['1']
            combo_pages.current(0)

def extract_table_from_image(image_path, min_conf=0, dist_thresh=25.0, min_size=2):
    # load the input image and convert it to grayscale
    image = cv2.imread(image_path)
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)

    # initialize a rectangular kernel that is ~5x wider than it is tall,
    # then smooth the image using a 3x3 Gaussian blur and then apply a
    # blackhat morph operator to find dark regions on a light background
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (51, 11))
    gray = cv2.GaussianBlur(gray, (3, 3), 0)
    blackhat = cv2.morphologyEx(gray, cv2.MORPH_BLACKHAT, kernel)

    # compute the Scharr gradient of the blackhat image and scale the
    # result into the range [0, 255]
    grad = cv2.Sobel(blackhat, ddepth=cv2.CV_32F, dx=1, dy=0, ksize=-1)
    grad = np.absolute(grad)
    (minVal, maxVal) = (np.min(grad), np.max(grad))
    grad = (grad - minVal) / (maxVal - minVal)
    grad = (grad * 255).astype("uint8")

    # apply a closing operation using the rectangular kernel to close
    # gaps in between characters, apply Otsu's thresholding method, and
    # finally a dilation operation to enlarge foreground regions
    grad = cv2.morphologyEx(grad, cv2.MORPH_CLOSE, kernel)
    thresh = cv2.threshold(grad, 0, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)[1]
    thresh = cv2.dilate(thresh, None, iterations=3)

    # find contours in the thresholded image and grab the largest one,
    # which we will assume is the stats table
    cnts, _ = cv2.findContours(thresh.copy(), cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    if len(cnts) == 0:
        return None, pytesseract.image_to_string(gray)
    tableCnt = max(cnts, key=cv2.contourArea)

    # compute the bounding box coordinates of the stats table and extract
    # the table from the input image
    (x, y, w, h) = cv2.boundingRect(tableCnt)
    table = image[y:y + h, x:x + w]

    # set the PSM mode to detect sparse text, and then localize text in
    # the table
    options = "--psm 6"
    results = pytesseract.image_to_data(cv2.cvtColor(table, cv2.COLOR_BGR2RGB), config=options, output_type=Output.DICT)

    # initialize a list to store the (x, y)-coordinates of the detected
    # text along with the OCR'd text itself
    coords = []
    ocrText = []
    full_text = ''

    # loop over each of the individual text localizations
    for i in range(0, len(results["text"])):
        # extract the bounding box coordinates of the text region from
        # the current result
        x = results["left"][i]
        y = results["top"][i]
        w = results["width"][i]
        h = results["height"][i]

        # extract the OCR text itself along with the confidence of the
        # text localization
        text = results["text"][i]
        conf = int(float(results["conf"][i]))

        # filter out weak confidence text localizations
        if conf > min_conf:
            coords.append((x, y, w, h))
            cleaned_text = text.strip()
            ocrText.append(cleaned_text)
            full_text += cleaned_text + ' '

    full_text = full_text.strip()

    if not coords:
        return None, full_text

    # extract all x-coordinates from the text prediction bounding boxes,
    # setting the y-coordinate value to zero
    xCoords = [(c[0], 0) for c in coords]

    # apply hierarchical agglomerative clustering to the coordinates
    clustering = AgglomerativeClustering(n_clusters=None, metric="manhattan", linkage="complete", distance_threshold=dist_thresh)
    clustering.fit(xCoords)

    # initialize our list of sorted clusters
    sortedClusters = []

    # loop over all clusters
    for l in np.unique(clustering.labels_):
        # extract the indexes for the coordinates belonging to the
        # current cluster
        idxs = np.where(clustering.labels_ == l)[0]

        # verify that the cluster is sufficiently large
        if len(idxs) >= min_size:
            # compute the average x-coordinate value of the cluster and
            # update our clusters list with the current label and the
            # average x-coordinate
            avg = np.average([coords[i][0] for i in idxs])
            sortedClusters.append((l, avg))

    # sort the clusters by their average x-coordinate and initialize our
    # data frame
    sortedClusters.sort(key=lambda x: x[1])
    df = pd.DataFrame()

    # loop over the clusters
    for (l, _) in sortedClusters:
        # extract the indexes for the coordinates belonging to the
        # current cluster
        idxs = np.where(clustering.labels_ == l)[0]

        # extract the y-coordinates from the elements in the current
        # cluster, then sort them from top-to-bottom
        yCoords = [coords[i][1] for i in idxs]
        sortedIdxs = sorted(range(len(yCoords)), key=lambda k: yCoords[k])

        # extract the OCR'd text for the current column
        cols = [ocrText[idxs[i]] for i in sortedIdxs]

        # construct a column for the current cluster (column) and append it to the dataframe
        # assuming the first entry is the header
        if cols:
            column = pd.DataFrame({cols[0]: cols[1:] if len(cols) > 1 else []})
            df = pd.concat([df, column], axis=1)

    return df, full_text

def process_file():
    input_file = entry_file_path.get()
    if not input_file:
        messagebox.showerror("Error", "Please select a file.")
        return

    page_range = combo_pages.get()
    output_type = combo_output.get()

    if not page_range or not output_type:
        messagebox.showerror("Error", "Please select page range and output type.")
        return

    base_name = os.path.basename(input_file).rsplit('.', 1)[0]
    desktop = os.path.join(os.path.expanduser("~"), "OneDrive", "Desktop")
    all_dfs = []
    all_text = ''

    if input_file.lower().endswith('.pdf'):
        num_pages = get_pdf_page_count(input_file)
        if page_range == 'All':
            start, end = 1, num_pages
        else:
            start, end = map(int, page_range.split('-'))
        page_imgs = convert_from_path(input_file, first_page=start, last_page=end)
        for idx, page_img in enumerate(page_imgs, start=start):
            temp_path = os.path.join(desktop, f"temp_page_{idx}.jpg")
            page_img.save(temp_path, 'JPEG')
            df, text = extract_table_from_image(temp_path)
            if df is not None and not df.empty:
                all_dfs.append((f"Page {idx}", df))
            all_text += text + '\n\n'
            os.remove(temp_path)
    else:
        df, text = extract_table_from_image(input_file)
        if df is not None and not df.empty:
            all_dfs.append(("Table 1", df))
        all_text = text

    has_tables = len(all_dfs) > 0

    ext_map = {'Excel': '.xlsx', 'TXT': '.txt', 'DOC': '.docx'}
    output_ext = ext_map[output_type]
    range_str = page_range.replace('-', '_') if page_range != 'All' else 'all'
    output = os.path.join(desktop, f"{base_name}_{range_str}{output_ext}")

    if output_type == 'Excel':
        writer = pd.ExcelWriter(output, engine='openpyxl')
        text_df = pd.DataFrame({'Extracted Text': all_text.split('\n')})
        text_df.to_excel(writer, sheet_name='Full Text', index=False)
        for sheet_name, df in all_dfs:
            df.to_excel(writer, sheet_name=sheet_name, index=False)
        writer.close()
    elif output_type == 'TXT':
        with open(output, 'w', encoding='utf-8') as f:
            f.write(all_text)
    elif output_type == 'DOC':
        doc = Document()
        doc.add_heading('Extracted Text', level=1)
        doc.add_paragraph(all_text)
        if has_tables:
            doc.add_heading('Extracted Tables', level=1)
            for sheet_name, df in all_dfs:
                doc.add_heading(sheet_name, level=2)
                table_doc = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
                table_doc.style = 'Table Grid'
                for j, col in enumerate(df.columns):
                    table_doc.cell(0, j).text = str(col)
                for row_idx, row in df.iterrows():
                    for j, val in enumerate(row):
                        table_doc.cell(row_idx + 1, j).text = str(val) if pd.notnull(val) else ''
        doc.save(output)

    messagebox.showinfo("Success", f"File saved to {output}")

def exit_app():
    root.quit()

# GUI
root = tk.Tk()
root.title("Data Extractor GUI")
root.geometry("500x300")

entry_file_path = tk.StringVar()

tk.Button(root, text="Select File", command=select_file).grid(row=0, column=0, columnspan=2, pady=10)

tk.Label(root, text="Select Page Range:").grid(row=1, column=0, pady=10)
combo_pages = ttk.Combobox(root, values=[], state="readonly")
combo_pages.grid(row=1, column=1, padx=10)

tk.Label(root, text="Select Output Type:").grid(row=2, column=0, pady=10)
combo_output = ttk.Combobox(root, values=['Excel', 'TXT', 'DOC'], state="readonly")
combo_output.grid(row=2, column=1, padx=10)
combo_output.current(0)

tk.Button(root, text="Process", command=process_file).grid(row=3, column=0, pady=20, padx=20)
tk.Button(root, text="Exit", command=exit_app).grid(row=3, column=1, pady=20)

root.mainloop()